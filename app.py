from flask import Flask, request, send_file, jsonify
from flask_cors import CORS
import os
import re
import sys
import logging
import tempfile
from datetime import datetime
from pathlib import Path
from typing import Dict, List, Tuple
from werkzeug.utils import secure_filename
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Configure detailed logging
logging.basicConfig(
    level=logging.DEBUG,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    handlers=[
        logging.StreamHandler(sys.stdout)
    ]
)
logger = logging.getLogger(__name__)

app = Flask(__name__)

# Use temporary directory for file operations
TEMP_DIR = tempfile.gettempdir()
ALLOWED_EXTENSIONS = {'docx'}

# Configure CORS
CORS(app, resources={
    r"/*": {
        "origins": "*",
        "methods": ["OPTIONS", "POST"],
        "allow_headers": ["Content-Type"]
    }
})

def allowed_file(filename: str) -> bool:
    """Check if the file extension is allowed."""
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def copy_cell_format(source_cell, target_cell):
    """Copy cell formatting from source to target."""
    try:
        # Copy paragraph formatting
        for src_para, tgt_para in zip(source_cell.paragraphs, target_cell.paragraphs):
            tgt_para.alignment = src_para.alignment
            for src_run, tgt_run in zip(src_para.runs, target_cell.paragraphs[0].runs):
                tgt_run.font.name = src_run.font.name
                tgt_run.font.size = src_run.font.size
                tgt_run.font.bold = src_run.font.bold
                tgt_run.font.italic = src_run.font.italic
    except Exception as e:
        logger.warning(f"Error copying cell format: {str(e)}")

def clean_course_title(title: str) -> str:
    """
    Clean course title according to specified rules.
    """
    if not title or 'Study Hall' in title:
        return ''
        
    # Store and remove '+' prefix (for AP/Honors courses)
    has_plus = title.startswith('+')
    title = title[1:].strip() if has_plus else title.strip()
    
    patterns = [
        # Grade indicators with section
        r'\bG\d+(?:-\d+)?(?=\s|$|\()',  # G10-2, G11-1
        r'\b(?:Grade\s*)?\d{1,2}(?:th)?\s*(?:Grade\s*)?(?:-\d+)?(?=\s|$)',  # 10, Grade 10, 10th Grade
        
        # Course group labels
        r'^(?:Senior|Junior)?\s*Electives\s*\d*\s*(?:\([^)]+\))?-',  # Senior Electives-, Electives 1 (G11)-
        r'^(?:Math|Science)\s+\d+[A-Z]?(?:-\d+)?-(?:\s*Group\s*\d+)?-?',  # Math 7A-2-, Science 8A-1-
        
        # Department prefixes
        r'^Career Planning\s*\d+(?:-\d+)?-?',
        r'^Visual Performing Arts\s*(?:G\d+(?:-\d+)?)?-',
        r'^Individual Society Environment\s*(?:G\d+(?:-\d+)?)?-',
        r'^Foreign Language-',
        r'^Military Training\s*(?:G\d+(?:-\d+)?)?-',
        r'^Social Studies-',
        
        # Group and section identifiers
        r'\s*Group\s*\d+\s*-',  # Group 1-
        r'-\d+(?:\s|$)',  # trailing -1, -2
        r'\s*\([^)]*\)',  # anything in parentheses
        
        # Clean up remaining patterns
        r'\s*-\s*(?=\S)',  # standalone hyphens with content after
        r'\s+-\s*$',  # trailing hyphens
    ]
    
    # Apply patterns in sequence
    for pattern in patterns:
        original = title
        title = re.sub(pattern, '', title, flags=re.IGNORECASE)
        if original != title:
            logger.debug(f"Pattern '{pattern}' changed '{original}' to '{title}'")
    
    # Clean up whitespace and multiple hyphens
    title = re.sub(r'\s+', ' ', title)
    title = re.sub(r'-+', '-', title)
    title = title.strip(' -')
    
    # Restore '+' prefix if it existed
    return f"+ {title}" if has_plus else title

def process_table(table) -> None:
    """
    Process table content with improved duplicate handling.
    """
    courses_by_semester = {'1st': [], '2nd': []}
    current_semester = None
    header_row = table.rows[0]  # Store header row
    
    # First pass: collect course data
    for row_index, row in enumerate(table.rows):
        cells = [cell.text.strip() for cell in row.cells]
        
        # Skip header row
        if row_index == 0:
            continue
            
        # Detect semester headers
        if len(cells) > 0:
            semester_text = ' '.join(cells).upper()
            if 'SEMESTER' in semester_text:
                if any(x in semester_text for x in ['1ST', 'FIRST']):
                    current_semester = '1st'
                elif any(x in semester_text for x in ['2ND', 'SECOND']):
                    current_semester = '2nd'
                logger.debug(f"Detected semester: {current_semester}")
                continue
        
        # Skip invalid rows
        if len(cells) < 3 or not current_semester:
            continue
            
        try:
            course_title = cells[0].strip()
            grade = cells[1].strip()
            gpa = cells[2].strip()
            
            # Skip rows with invalid data
            if not course_title or not grade.replace('.', '').isdigit():
                continue
                
            # Clean course title
            clean_title = clean_course_title(course_title)
            if not clean_title:  # Skip empty or removed courses
                logger.debug(f"Skipping removed course: {course_title}")
                continue
                
            # Store course data with original formatting
            courses_by_semester[current_semester].append({
                'title': clean_title,
                'grade': grade,
                'gpa': gpa,
                'original_row': row
            })
            logger.debug(f"Added course: {clean_title} (Grade: {grade}, GPA: {gpa})")
            
        except Exception as e:
            logger.error(f"Error processing row: {str(e)}")
            continue
    
    # Second pass: handle duplicates
    for semester in courses_by_semester:
        # Group courses by cleaned title
        courses_by_title = {}
        for course in courses_by_semester[semester]:
            title = course['title']
            if title not in courses_by_title:
                courses_by_title[title] = []
            courses_by_title[title].append(course)
        
        # Process each group of courses
        unique_courses = []
        for title, course_group in courses_by_title.items():
            if len(course_group) == 1:
                unique_courses.append(course_group[0])
                continue
                
            # Group by grade and GPA
            seen = set()
            for course in course_group:
                key = (course['grade'], course['gpa'])
                if key not in seen:
                    unique_courses.append(course)
                    seen.add(key)
                else:
                    logger.debug(f"Removing duplicate: {title} {key}")
        
        courses_by_semester[semester] = unique_courses
    
    # Rebuild table
    try:
        # Clear existing content except header
        while len(table.rows) > 1:
            table._element.remove(table.rows[-1]._element)
        
        # Add processed courses
        for semester, courses in courses_by_semester.items():
            if not courses:
                continue
                
            # Add semester header
            semester_row = table.add_row()
            semester_cell = semester_row.cells[0]
            semester_cell.text = f"{semester} Semester"
            if len(semester_cell.paragraphs) > 0:
                semester_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add courses
            for course in courses:
                new_row = table.add_row()
                
                # Copy formatting from original row if available
                if 'original_row' in course:
                    for i, (new_cell, old_cell) in enumerate(zip(new_row.cells, course['original_row'].cells)):
                        copy_cell_format(old_cell, new_cell)
                
                # Set content
                new_row.cells[0].text = course['title']
                new_row.cells[1].text = course['grade']
                new_row.cells[2].text = course['gpa']
                
                # Center align grade and GPA cells
                for cell in new_row.cells[1:]:
                    if len(cell.paragraphs) > 0:
                        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
    except Exception as e:
        logger.error(f"Error rebuilding table: {str(e)}")
        raise

def process_report_card(filepath: str) -> str:
    """Process the report card document while preserving formatting."""
    try:
        logger.debug(f"Opening document: {filepath}")
        doc = Document(filepath)
        
        # Process each table in the document
        for table_index, table in enumerate(doc.tables):
            logger.debug(f"Processing table {table_index}")
            process_table(table)
        
        # Save processed document to temporary file
        output_path = os.path.join(TEMP_DIR, f"processed_{os.path.basename(filepath)}")
        logger.debug(f"Saving processed document to: {output_path}")
        doc.save(output_path)
        return output_path
        
    except Exception as e:
        logger.error(f"Error processing report card: {str(e)}")
        raise

@app.route("/upload", methods=["POST"])
def upload_file():
    """Handle file upload and process the report card."""
    temp_dir = Path(tempfile.gettempdir())
    input_file = None
    output_file = None
    
    try:
        if "file" not in request.files:
            return jsonify({"error": "No file part"}), 400
        
        file = request.files["file"]
        if file.filename == "":
            return jsonify({"error": "No selected file"}), 400
        
        if not allowed_file(file.filename):
            return jsonify({"error": "Invalid file type. Please upload a .docx file"}), 400
        
        # Create unique filenames
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        input_file = temp_dir / f"input_{timestamp}_{secure_filename(file.filename)}"
        
        logger.debug(f"Saving input file to {input_file}")
        file.save(str(input_file))
        
        logger.debug("Processing report card")
        processed_filepath = process_report_card(str(input_file))
        
        logger.debug(f"Sending processed file: {processed_filepath}")
        return send_file(
            processed_filepath,
            as_attachment=True,
            download_name=f"processed_{file.filename}",
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        
    except Exception as e:
        logger.error(f"Error in upload_file: {str(e)}")
        return jsonify({"error": str(e)}), 500
        
    finally:
        # Clean up temporary files
        for file in [input_file, output_file]:
            if file and Path(str(file)).exists():
                try:
                    Path(str(file)).unlink()
                    logger.debug(f"Cleaned up temporary file: {file}")
                except Exception as e:
                    logger.error(f"Error removing temporary file {file}: {str(e)}")

if __name__ == "__main__":
    app.run(host="0.0.0.0", port=5000)